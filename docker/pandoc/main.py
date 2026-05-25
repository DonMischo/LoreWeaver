"""
Foliantica Pandoc export service.

POST /convert
  body (JSON):
    html:             str          — HTML content of the book
    format:           "pdf" | "epub" | "docx"
    title:            str
    author:           str | None
    language:         str | None   — BCP-47 code, e.g. "en", "de"
    cover:            str | None   — base64-encoded PNG/JPEG for EPUB cover
    font:             str | None   — body font name
    heading_font:     str | None   — separate heading font (EPUB CSS only for now)
    heading_align:    str          — "center" | "left"
    h1_size:          str          — CSS size, e.g. "2em"
    h2_size:          str
    h3_size:          str
    h3_style:         str          — "italic" | "normal" | "bold"
    paragraph_indent: str          — e.g. "1.5em" or "0"
    text_align:       str          — "justify" | "left"
    pdf_margin:       str          — e.g. "2.5cm" or "1in"
    page_numbers:     bool
    line_spacing:     str          — "1" | "1.5" | "2"
    font_size:        str          — "10pt" | "11pt" | "12pt"

GET /fonts
  Returns { fonts: [str] } — list of font families available to fontconfig

GET /health
  Health check
"""

import base64
import os
import re as _re
import subprocess
import tempfile
import urllib.request as _urlreq
from pathlib import Path
from fastapi import FastAPI, HTTPException
from fastapi.responses import Response
from pydantic import BaseModel

app = FastAPI(title="Foliantica Pandoc")

# System font directory — writable in the container, auto-found by fontconfig
_FONT_INSTALL_DIR = Path("/usr/local/share/fonts/gfonts")


# ── Google Font helper ────────────────────────────────────────────────────────

def _ensure_google_font(name: str) -> bool:
    """Download a Google Font TTF into the system fonts dir and register it.
    Returns True if the font is (now) available, False on failure."""
    safe     = name.replace(" ", "_")
    dest_dir = _FONT_INSTALL_DIR / safe

    # Already cached — skip download but check fontconfig registration
    if any(dest_dir.glob("*.ttf")) or any(dest_dir.glob("*.otf")):
        return True

    # Use an old User-Agent so Google Fonts returns TTF format
    css_url = f"https://fonts.googleapis.com/css?family={name.replace(' ', '+')}"
    req = _urlreq.Request(css_url, headers={
        "User-Agent": "Mozilla/4.0 (compatible; MSIE 5.5; Windows NT 4.0)"
    })
    try:
        with _urlreq.urlopen(req, timeout=15) as r:
            css = r.read().decode()
    except Exception:
        return False

    urls = _re.findall(
        r"url\((https://fonts\.gstatic\.com/[^)]+\.(?:ttf|otf))\)", css
    )
    if not urls:
        return False

    dest_dir.mkdir(parents=True, exist_ok=True)
    downloaded = 0
    for i, url in enumerate(urls[:4]):   # up to 4 weight variants
        ext = url.rsplit(".", 1)[-1]
        try:
            with _urlreq.urlopen(url, timeout=20) as r:
                (dest_dir / f"{safe}_{i}.{ext}").write_bytes(r.read())
            downloaded += 1
        except Exception:
            pass

    if downloaded == 0:
        return False

    # Rebuild fontconfig cache so xelatex and fc-list can see the new fonts
    subprocess.run(["fc-cache", "-f", str(_FONT_INSTALL_DIR)],
                   capture_output=True, timeout=15)
    return True


def _font_available(name: str) -> bool:
    """Check if a font family is available to fontconfig."""
    result = subprocess.run(
        ["fc-list", f":family={name}"],
        capture_output=True, text=True, timeout=5,
    )
    return bool(result.stdout.strip())


def _resolve_font(name: str | None) -> str | None:
    """Ensure font is available locally; try Google Fonts if not. Returns name or None."""
    if not name:
        return None
    if _font_available(name):
        return name
    if _ensure_google_font(name):
        return name
    return None  # could not resolve — fall back to default


# ── DOCX reference document builder ──────────────────────────────────────────

def _build_reference_docx(dest: Path, font: str, line_spacing: str,
                           margin: str, paragraph_indent: str,
                           heading_align: str, author: str, title: str) -> None:
    """
    Build a reference.docx via python-docx that encodes:
      - Body font (Normal style), font size, line spacing, first-line indent
      - Page margins
      - Running header: AUTHOR / TITLE / Page #
    Pandoc then uses it as a style template via --reference-doc.
    """
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import re as _r

    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────────────
    def _to_inches(val: str) -> float:
        """Convert margin string (e.g. '1in', '2.5cm', '2em') to inches."""
        val = val.strip()
        if val.endswith("in"):
            return float(val[:-2])
        if val.endswith("cm"):
            return float(val[:-2]) / 2.54
        if val.endswith("mm"):
            return float(val[:-2]) / 25.4
        # em / other — treat 1em ≈ 0.5in
        m = _r.match(r"([\d.]+)", val)
        return float(m.group(1)) * 0.5 if m else 1.0

    margin_in = _to_inches(margin)
    for section in doc.sections:
        section.top_margin    = Inches(margin_in)
        section.bottom_margin = Inches(margin_in)
        section.left_margin   = Inches(margin_in)
        section.right_margin  = Inches(margin_in)

    # ── Body font (Normal style) ──────────────────────────────────────────────
    _LS_MAP = {"1": WD_LINE_SPACING.SINGLE, "1.5": WD_LINE_SPACING.ONE_POINT_FIVE,
               "2": WD_LINE_SPACING.DOUBLE}
    ls = _LS_MAP.get(line_spacing, WD_LINE_SPACING.DOUBLE)

    _indent_map = {"0": 0.0, "1.5em": 0.5, "1em": 0.33}
    indent_in = _indent_map.get(paragraph_indent, 0.5)

    normal = doc.styles["Normal"]
    nf = normal.font
    nf.name = font
    nf.size = Pt(12)
    npf = normal.paragraph_format
    npf.line_spacing_rule = ls
    npf.first_line_indent = Inches(indent_in) if indent_in else None
    npf.space_after = Pt(0)

    # ── Heading styles ────────────────────────────────────────────────────────
    halign = WD_ALIGN_PARAGRAPH.CENTER if heading_align == "center" else WD_ALIGN_PARAGRAPH.LEFT
    for h_name in ("Heading 1", "Heading 2", "Heading 3"):
        try:
            hs = doc.styles[h_name]
            hs.font.name  = font
            hs.font.bold  = True
            hs.paragraph_format.alignment     = halign
            hs.paragraph_format.space_before  = Pt(24)
            hs.paragraph_format.space_after   = Pt(6)
            hs.paragraph_format.first_line_indent = None
        except Exception:
            pass

    # ── Running header: AUTHOR / TITLE / Page # ───────────────────────────────
    if author or title:
        from docx.oxml.shared import OxmlElement
        section = doc.sections[0]
        section.different_first_page_header = False
        header = section.header
        if not header.paragraphs:
            header.add_paragraph()
        hp = header.paragraphs[0]
        hp.clear()
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hp.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

        # Build: "AUTHOR / TITLE / {page #}"
        parts: list[str] = []
        if author:
            parts.append(author.split()[-1].upper())   # last name
        if title:
            short = title[:30].upper()
            parts.append(short)

        header_text = " / ".join(parts) + " / " if parts else ""
        if header_text:
            run = hp.add_run(header_text)
            run.font.name  = font
            run.font.size  = Pt(12)

        # Insert PAGE field
        fld_char1 = OxmlElement("w:fldChar")
        fld_char1.set(qn("w:fldCharType"), "begin")
        instr_text = OxmlElement("w:instrText")
        instr_text.text = " PAGE "
        fld_char2 = OxmlElement("w:fldChar")
        fld_char2.set(qn("w:fldCharType"), "end")
        pg_run = hp.add_run()
        pg_run.font.name = font
        pg_run.font.size = Pt(12)
        pg_run._r.append(fld_char1)
        pg_run._r.append(instr_text)
        pg_run._r.append(fld_char2)

    doc.save(str(dest))


# ── Request model ─────────────────────────────────────────────────────────────

class ConvertRequest(BaseModel):
    html:             str
    format:           str                  # "pdf" | "epub" | "docx"
    title:            str    = "Untitled"
    author:           str | None = None
    language:         str | None = "en"
    cover:            str | None = None    # base64 image for EPUB cover
    # Typography
    font:             str | None = None
    heading_font:     str | None = None
    heading_align:    str        = "center"
    h1_size:          str        = "2em"
    h2_size:          str        = "1.5em"
    h3_size:          str        = "1.25em"
    h3_style:         str        = "italic"
    paragraph_indent: str        = "1.5em"
    text_align:       str        = "justify"
    pdf_margin:       str        = "2.5cm"
    page_numbers:     bool       = True
    line_spacing:     str        = "1.5"
    font_size:        str        = "12pt"


# ── Endpoints ─────────────────────────────────────────────────────────────────

@app.get("/health")
def health():
    return {"status": "ok"}


@app.get("/fonts")
def list_fonts():
    """Return all font families available to fontconfig / XeLaTeX."""
    result = subprocess.run(
        ["fc-list", ":", "family"],
        capture_output=True, text=True, timeout=10,
    )
    fonts: set[str] = set()
    for line in result.stdout.splitlines():
        # fc-list may output "Family,Variant Style" — take first part
        name = line.split(",")[0].strip()
        if name:
            fonts.add(name)
    return {"fonts": sorted(fonts)}


_LS_TO_STRETCH = {"1": "1.0", "1.5": "1.5", "2": "2.0"}
_FS_TO_PT      = {"10pt": "10pt", "11pt": "11pt", "12pt": "12pt"}


@app.post("/convert")
def convert(req: ConvertRequest):
    if req.format not in ("pdf", "epub", "docx"):
        raise HTTPException(400, "format must be 'pdf', 'epub', or 'docx'")

    with tempfile.TemporaryDirectory() as tmp:
        tmp  = Path(tmp)
        src  = tmp / "input.html"
        out  = tmp / f"output.{req.format}"

        src.write_text(req.html, encoding="utf-8")

        cmd = [
            "pandoc", str(src), "-o", str(out),
            "--metadata", f"title={req.title}",
            "--standalone",
        ]

        if req.author:
            cmd += ["--metadata", f"author={req.author}"]
        if req.language:
            cmd += ["--metadata", f"lang={req.language}"]

        if req.format == "pdf":
            # Resolve fonts (download from Google Fonts if needed).
            # NOTE: fonts are installed to /usr/local/share/fonts — no FONTCONFIG_PATH
            # override needed. xelatex finds them via the standard fontconfig search.
            body_font = _resolve_font(req.font)

            pdf_vars = [
                ("geometry",    f"margin={req.pdf_margin}"),
                ("fontsize",    _FS_TO_PT.get(req.font_size, "12pt")),
                ("linestretch", _LS_TO_STRETCH.get(req.line_spacing, "1.5")),
            ]
            if body_font:
                pdf_vars.append(("mainfont", body_font))
            if not req.page_numbers:
                pdf_vars.append(("pagestyle", "empty"))

            for k, v in pdf_vars:
                cmd += ["-V", f"{k}:{v}"]
            cmd += ["--pdf-engine=xelatex"]
            env = None  # use ambient environment — fonts are in standard system path

        elif req.format == "docx":
            # Build a reference.docx that encodes font, spacing, margins, and header
            body_font     = req.font or "Times New Roman"
            ref_docx_path = tmp / "reference.docx"
            try:
                _build_reference_docx(
                    dest             = ref_docx_path,
                    font             = body_font,
                    line_spacing     = req.line_spacing,
                    margin           = req.pdf_margin,
                    paragraph_indent = req.paragraph_indent,
                    heading_align    = req.heading_align,
                    author           = req.author or "",
                    title            = req.title,
                )
                cmd += [f"--reference-doc={ref_docx_path}"]
            except Exception as e:
                # If python-docx is unavailable or fails, proceed without reference doc
                pass
            env = None

        elif req.format == "epub":
            cmd += ["--epub-chapter-level=2"]
            if req.cover:
                cover_path = tmp / "cover.jpg"
                cover_path.write_bytes(base64.b64decode(req.cover))
                cmd += [f"--epub-cover-image={cover_path}"]

            # Generate an inline CSS for the EPUB
            body_font    = req.font or "Georgia"
            h_font       = req.heading_font or body_font
            indent_css   = req.paragraph_indent if req.paragraph_indent != "0" else "0"
            h3_style_css = {
                "italic": "font-style: italic;",
                "bold":   "font-weight: bold;",
                "normal": "",
            }.get(req.h3_style, "font-style: italic;")

            epub_css = f"""\
body {{
  font-family: "{body_font}", Georgia, serif;
  font-size: 1em;
  line-height: {_LS_TO_STRETCH.get(req.line_spacing, "1.5")};
  text-align: {req.text_align};
  margin: 2em;
}}
h1, h2, h3, h4 {{
  font-family: "{h_font}", Georgia, serif;
  text-align: {req.heading_align};
}}
h1 {{ font-size: {req.h1_size}; }}
h2 {{ font-size: {req.h2_size}; }}
h3 {{ font-size: {req.h3_size}; {h3_style_css} }}
p {{ margin: 0; text-indent: {indent_css}; }}
p:first-of-type, h1 + p, h2 + p, h3 + p {{ text-indent: 0; }}
figure {{ margin: 1.5em auto; text-align: center; }}
figure img {{ max-width: 100%; height: auto; }}
figcaption {{ font-size: 0.85em; opacity: 0.7; margin-top: 0.4em; }}
"""
            css_path = tmp / "style.css"
            css_path.write_text(epub_css, encoding="utf-8")
            cmd += [f"--css={css_path}"]
            env = None

        result = subprocess.run(
            cmd,
            capture_output=True,
            text=False,
            timeout=120,
            env=env,
        )

        if result.returncode != 0:
            error = result.stderr.decode("utf-8", errors="replace")
            raise HTTPException(500, f"Pandoc error: {error[:500]}")

        data = out.read_bytes()

    media_types = {
        "pdf":  "application/pdf",
        "epub": "application/epub+zip",
        "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    }
    return Response(content=data, media_type=media_types[req.format])
